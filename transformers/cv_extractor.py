"""Extração de texto de arquivos PDF e DOCX para parsing de currículo."""

from __future__ import annotations
from typing import BinaryIO


def extrair_texto_pdf(file: BinaryIO) -> str:
    import pdfplumber

    linhas = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            texto = page.extract_text(layout=True)
            if texto:
                linhas.append(texto)
    return "\n".join(linhas)


def extrair_texto_docx(file: BinaryIO) -> str:
    from docx import Document

    doc = Document(file)
    partes = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            partes.append(t)
    for table in doc.tables:
        for row in table.rows:
            celulas = [c.text.strip() for c in row.cells if c.text.strip()]
            if celulas:
                partes.append("  ".join(celulas))
    return "\n".join(partes)


def extrair_texto(file: BinaryIO, nome_arquivo: str) -> str:
    """Detecta tipo pelo nome e extrai texto do arquivo."""
    ext = nome_arquivo.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extrair_texto_pdf(file)
    elif ext in ("docx", "doc"):
        return extrair_texto_docx(file)
    raise ValueError(f"Formato não suportado: .{ext}")
